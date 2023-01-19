from django.shortcuts import render
import urllib.request
import json
import io
from docx import Document
from docx.shared import Inches
import requests

# Create your views here.
def index(request):




    if request.method == "POST":
        pokemon = request.POST['pokemon'].lower()
        pokemon = pokemon.replace(' ', '%20')
        url_pokeapi = urllib.request.Request(f'https://pokeapi.co/api/v2/pokemon/{pokemon}/')
        url_pokeapi.add_header('User-Agent', 'blastoise')

        source = urllib.request.urlopen(url_pokeapi).read()

        list_of_data = json.loads(source)

        altura = list_of_data['height']/10
        peso = list_of_data['weight']/10

        data = {
            "id": str(list_of_data['id']),
            "nome": str(list_of_data['species']['name']).capitalize(),
            "altura": str(altura),
            "peso": str(peso),
            "foto": str(list_of_data['sprites']['other']['official-artwork']['front_default']),
            "foto-shiny": str(list_of_data['sprites']['other']['official-artwork']['front_shiny'])
        }

        response = requests.get(data['foto'], stream=True)
        foto = io.BytesIO(response.content)
        response = requests.get(data['foto-shiny'], stream=True)
        foto_shiny = io.BytesIO(response.content)

        document = Document()
        document.add_heading(f"{data['nome']} - {data['id']}", 0)
        document.add_picture(foto, width=Inches(1.25))
        document.add_picture(foto_shiny, width=Inches(1.25))
        
        document.save('demo.docx')

        return render(request, 'index.html' ,data , )

    else:
        data = {}

        return render(request, 'index.html' ,data )